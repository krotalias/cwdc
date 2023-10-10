#!/usr/bin/env python
# coding: UTF-8
#
## @package _03d_pascal
#  Showing a range of rows of a Pascal triangle.
#   <pre>
#   0       1                                 C(0,0)
#   1      1 1                            C(1,0)  C(1,1)
#   2     1 2 1                        C(2,0)  C(2,1)  C(2,2)
#   3    1 3 3 1                   C(3,0)  C(3,1)  C(3,2)  C(3,3)
#   4   1 4 6 4 1              C(4,0)  C(4,1)  C(4,2)  C(4,3)  C(4,4)
#   5  1 5 10 10 5 1       C(5,0)  C(5,1)  C(5,2)  C(5,3)  C(5,4)  C(5,5)
#
#      0 1 2  3  4 5
#  </pre>
#
#   @author Paulo Roma
#   @since 13/04/2009
#   @see http://en.wikipedia.org/wiki/Pascal_triangle

import sys

from builtins import input
from math import factorial
try:
    from scipy.special import comb
except:
    from scipy.special import binom as comb

## Compute n lines of the binomial array.
#  Complexity is @f$1 + 2\ +\ ...\ +\ n@f$ elements.
#  This is arithmetic progression that sums to @f${n*(n+1)}\over{2}@f$,
#  which is in @f$O(n^2)@f$.
#   <PRE>
#   0 - 1
#   1 - 1 1
#   2 - 1 2 1
#   3 - 1 3 3 1
#   4 - 1 4 6 4 1
#   5 - 1 5 10 10 5 1
#   </PRE>
#
#  @param n number of lines.
#  @return binomial array as a list of lines.
#
def binomial_array(n):
    bc = []  # table of binomial coefficients
    for i in range(n + 1):
        bc.append([0] * (i + 1))
        bc[i][0] = 1
        bc[i][i] = 1
    for i in range(1, n + 1):
        for j in range(1, i):
            bc[i][j] = bc[i - 1][j - 1] + bc[i - 1][j]
    return bc

## Compute @f$n@f$ choose @f$m@f$ or Combination @f$(n,m)@f$.
#
#  @param n number of objects to choose from.
#  @param m number of places.
#  @param method how to get the result.
#  @return @f$C{{n}\choose{m}} = {{n!} \over {(n-m)!\ m!}} = {{n (n-1) ... (n-m+1)} \over {m!}}@f$, number of ways of choosing m objects from n.
#  @see https://en.wikipedia.org/wiki/Combination
#  @see https://docs.scipy.org/doc/scipy/reference/tutorial/special.html
#  @see https://docs.scipy.org/doc/scipy/reference/generated/scipy.special.comb.html#scipy.special.comb
#
def binomial_coefficient(n, m, method=2):
    if m > n:
        return 0

    if method == 0:
        return binomial_array(n)[n][m]
    elif method == 1:
        return factorial(n) // (factorial(m) * factorial(n - m))
    else:
        return comb(n, m, exact=True)

## Compute the probability of getting at most @f$x@f$ successes in @f$n@f$ trials,
#  which is the sum of the first @f$x@f$ terms of the binomial expansion @f$(p+q)^n@f$.
#
#  In other words, the probability of obtaining at most @f$x@f$ successes in @f$n@f$ independent trials,
#  each of which has a probability @f$p@f$ of success, and @f$q = (1-p)@f$ of failure.
#
#  That is, if @f$X@f$ denotes the number of successes, it returns:
#  - @f$P(X \le x) = \sum_{i=0}^{x} {C {{n}\choose{i}}\ p^i\ (1-p)^{n-i}}@f$.
#
#  @param p probability of success.
#  @param n number of trials.
#  @param x number of successes.
#  @return the probability of getting at most x successes in n independent trials.
#  @see https://www.khanacademy.org/math/ap-statistics/probability-ap/probability-multiplication-rule/v/getting-at-least-one-heads
#  @see http://www.wolframalpha.com/widgets/view.jsp?id=d821210668c6cc5a02db1069cc52464f
#  @see https://mat.iitm.ac.in/home/vetri/public_html/statistics/binomial.pdf
#  @see https://www.khanacademy.org/math/statistics-probability/random-variables-stats-library/binomial-random-variables/v/visualizing-a-binomial-distribution
#
def binomial_cumulative_distribution(p, n, x=None):
    if n == 0:
        return 1
    if x is None:
        x = n
    if x > n:
        return 0
    if p < 0 or p > 1:
        return 0
    q = 1.0 - p
    value = 0
    ban = binomial_array(n)[n]  # get the nth row of the array
    for i in range(0, x + 1):
        value += ban[i] * p**(n - i) * q**i
    return value


## dictionary for mapping characters to superscript.
superscript_map = {
    "0": "⁰", "1": "¹", "2": "²", "3": "³", "4": "⁴", "5": "⁵", "6": "⁶",
    "7": "⁷", "8": "⁸", "9": "⁹", "a": "ᵃ", "b": "ᵇ", "c": "ᶜ", "d": "ᵈ",
    "e": "ᵉ", "f": "ᶠ", "g": "ᵍ", "h": "ʰ", "i": "ᶦ", "j": "ʲ", "k": "ᵏ",
    "l": "ˡ", "m": "ᵐ", "n": "ⁿ", "o": "ᵒ", "p": "ᵖ", "q": "۹", "r": "ʳ",
    "s": "ˢ", "t": "ᵗ", "u": "ᵘ", "v": "ᵛ", "w": "ʷ", "x": "ˣ", "y": "ʸ",
    "z": "ᶻ", "A": "ᴬ", "B": "ᴮ", "C": "ᶜ", "D": "ᴰ", "E": "ᴱ", "F": "ᶠ",
    "G": "ᴳ", "H": "ᴴ", "I": "ᴵ", "J": "ᴶ", "K": "ᴷ", "L": "ᴸ", "M": "ᴹ",
    "N": "ᴺ", "O": "ᴼ", "P": "ᴾ", "Q": "Q", "R": "ᴿ", "S": "ˢ", "T": "ᵀ",
    "U": "ᵁ", "V": "ⱽ", "W": "ᵂ", "X": "ˣ", "Y": "ʸ", "Z": "ᶻ", "+": "⁺",
    "-": "⁻", "=": "⁼", "(": "⁽", ")": "⁾"}

subscript_map = {
    "0": "₀", "1": "₁", "2": "₂", "3": "₃", "4": "₄", "5": "₅", "6": "₆",
    "7": "₇", "8": "₈", "9": "₉", "a": "ₐ", "b": "♭", "c": "꜀", "d": "ᑯ",
    "e": "ₑ", "f": "բ", "g": "₉", "h": "ₕ", "i": "ᵢ", "j": "ⱼ", "k": "ₖ",
    "l": "ₗ", "m": "ₘ", "n": "ₙ", "o": "ₒ", "p": "ₚ", "q": "૧", "r": "ᵣ",
    "s": "ₛ", "t": "ₜ", "u": "ᵤ", "v": "ᵥ", "w": "w", "x": "ₓ", "y": "ᵧ",
    "z": "₂", "A": "ₐ", "B": "₈", "C": "C", "D": "D", "E": "ₑ", "F": "բ",
    "G": "G", "H": "ₕ", "I": "ᵢ", "J": "ⱼ", "K": "ₖ", "L": "ₗ", "M": "ₘ",
    "N": "ₙ", "O": "ₒ", "P": "ₚ", "Q": "Q", "R": "ᵣ", "S": "ₛ", "T": "ₜ",
    "U": "ᵤ", "V": "ᵥ", "W": "w", "X": "ₓ", "Y": "ᵧ", "Z": "Z", "+": "₊",
    "-": "₋", "=": "₌", "(": "₍", ")": "₎"}

## character translation table.
try:
    trans = str.maketrans(
        ''.join(superscript_map.keys()),
        ''.join(superscript_map.values()))

    sub_trans = str.maketrans(
        ''.join(subscript_map.keys()),
        ''.join(subscript_map.values()))
except:
    trans = {ord(c): ord(t) for c, t in zip(unicode(''.join(superscript_map.keys()), 'utf-8'),
                                            unicode(''.join(superscript_map.values()), 'utf-8'))}

    sub_trans = {ord(c): ord(t) for c, t in zip(unicode(''.join(subscript_map.keys()), 'utf-8'),
                                                unicode(''.join(subscript_map.values()), 'utf-8'))}

## Return a superscript string for the given value.
#
#  @param val a numeric string.
#  @param type select the superscript method.
#  @return a new superscript string.
#
def exponent(val, type=False):
    if type:
        return '<sup>' + val + '</sup>'
    elif trans is not None:
        if isinstance(val, bytes):
            val = unicode(val, 'utf-8')
        return val.translate(trans)
    else:
        return '^' + val

## Compute the binomial expansion of @f$(x+y)^n@f$.
#
#  For e = 5, return: <br>
#  - @f$x^5 + 5x^4y + 10x^3y^2 + 10x^2y^3 + 5xy^4 + y^5@f$ (html)
#  - x⁵ + 5x⁴y + 10x³y² + 10x²y³ + 5xy⁴ + y⁵ (unicode translation table)
#  - x^5 + 5x^4y + 10x^3y^2 + 10x^2y^3 + 5xy^4 + y^5 (None)
#
#  @param e exponent n.
#  @return a string representing the expansion.
#  @see https://www.khanacademy.org/math/precalculus/prob-comb/prob-combinatorics-precalc/v/exactly-three-heads-in-five-flips
#  @see https://www.khanacademy.org/math/precalculus/prob-comb/combinations/v/combination-formula
#  @see https://www.khanacademy.org/math/statistics-probability/random-variables-stats-library/binomial-random-variables/v/binomial-distribution
#
def binomial_expansion(e):
    if e == 0:
        return '1'
    bstr = ''
    s = str(e)
    bstr = 'x'
    if s != '1':
        bstr += exponent(s)
    bstr += ' + '
    ba = binomial_array(e)
    for i in range(1, e):
        sx = str(e - i)
        sy = str(i)
        b = str(ba[e][i])
        bstr += b + 'x'
        if sx != '1':
            bstr += exponent(sx)
        bstr += 'y'
        if sy != '1':
            bstr += exponent(sy)
        bstr += ' + '
    bstr += 'y'
    if s != '1':
        bstr += exponent(s)
    return bstr

##
#   Computes the next row of a Pascal triangle, given the current row.
#
#   @param curr_row current row.
#   @return next_row: list(map(sum, zip([0] + curr_row, curr_row + [0])))
#
def getNextRow(curr_row):

    next_row = [1]                                   # 1 on the left side
    for i in range(1, len(curr_row)):
        next_row += [curr_row[i - 1] + curr_row[i]]  # sum of parents
    next_row += [1]

    return next_row

##
#   Formats a row of the Pascal triangle.
#
#   @param last_level  last_level.
#   @param r           row to be printed.
#   @param pad         whether to pad the row with blanks to the left.
#   @return string representing the row.
#
def fmtRow(last_level, r, pad=True):
    # adds some blanks for getting the triangle shape
    txt = ' ' * (last_level - len(r) + 2) if pad else ''
    txt += ' '.join(map(str, r))                    # show it
    return txt.rstrip()

##
#   Creates the pascal triangle between two levels.
#
#   @param first_level first level to be printed.
#   @param last_level  height of the Pascal triangle.
#   @param center whether to center the lines.
#   @see https://www.cut-the-knot.org/arithmetic/combinatorics/PascalTriangleProperties.shtml
#   @see https://www.johndcook.com/blog/2016/07/05/distribution-of-numbers-in-pascals-triangle/
#
def pascal(first_level, last_level, center=True):
    # the sum of numbers in each line is 2**n
    ndig = len(str(2**last_level)) + 1
    # therefore, it is an upper estimate for the number of digits
    maxlen = (last_level + 1) * ndig
    if last_level > 11:
        maxlen -= 4 * (ndig + 1)

    row = [1]                                       # Top of Pascal triangle
    # level begins at "0"
    for l in range(0, last_level + 1):
        # l is now in the range [first_level, last_level].
        if l >= first_level:
            if center:
                print(fmtRow(last_level, row, False).center(maxlen))
            else:
                print(fmtRow(last_level, row))
        row = getNextRow(row)                       # new row

##
#   Creates the pascal triangle between two levels, using the binomial array.
#   The lines are centered in respect to the last line of the triangle.
#   The output is written to a docx file.
#
#   @param first_level first level to be printed.
#   @param last_level  height of the Pascal triangle.
#   @param center whether to center the lines.
#
def pascal2(first_level, last_level, center=True):
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx import Document
    except ModuleNotFoundError as e:
        print(e)
        return

    document = Document()
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ba = binomial_array(last_level)

    # get length of last line
    lenmax = len(' '.join(map(str, ba[last_level])))
    for i in range(first_level, last_level + 1):
        run = paragraph.add_run(fmtRow(last_level, ba[i], False) + '\n')
        run.style = 'Emphasis'
        if center:
            print(fmtRow(last_level, ba[i], False).center(lenmax))
        else:
            print(fmtRow(last_level, ba[i]))
    print('')

    document.save('Pascal.docx')
    for p in document.paragraphs:
        print(p.text)

##
#   Main function for testing.
#
#   @param argv first and last levels of the Pascal triangle.
#
def main(argv=None):
    if argv is None:
        argv = sys.argv

    try:
        if (len(argv) > 2):
            first_level = int(argv[1])
            last_level = int(argv[2])
        else:
            a = input(
                "Enter first and last levels of Pascal triangle, separated by ',': ").split(',')
            first_level, last_level = int(a[0]), int(a[1])
    except Exception as e:
        print("Invalid values. %s. Assuming 3,12" % e)
        first_level = 3
        last_level = 12

    pascal(first_level, last_level, False)

    print("\nThe mean of the coefficients of the last row is: %f" %
          (2**last_level / float(last_level + 1)))

    pascal2(first_level, last_level)

    for i in range(last_level + 1):
        print(binomial_expansion(i))

    # in general, head is associated with success and tail to fail.
    heads = first_level
    flips = last_level
    print("Probability of getting at most %d heads in %d flips is: %f\n" %
          (heads, flips, binomial_cumulative_distribution(0.5, flips, heads)))


if __name__ == "__main__":
    sys.exit(main())
